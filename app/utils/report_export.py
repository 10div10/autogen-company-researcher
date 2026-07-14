"""Converts a markdown-ish research report string into downloadable
PDF and DOCX files. Deliberately simple/regex-based parsing -- good
enough for the headings/bullets the Writer agent produces, no heavy
markdown-parsing dependency required."""

import os
import re
from docx import Document
from docx.shared import Pt
from fpdf import FPDF

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "outputs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def _safe_filename(company_name: str, ext: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", company_name.strip()).strip("_").lower()
    return f"{slug or 'company'}_research_report.{ext}"


def _parse_lines(markdown_text: str):
    """Yields (kind, text) tuples: kind in {'h1','h2','h3','bullet','text'}."""
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            yield "h3", line[4:].strip()
        elif line.startswith("## "):
            yield "h2", line[3:].strip()
        elif line.startswith("# "):
            yield "h1", line[2:].strip()
        elif re.match(r"^[-*]\s+", line):
            yield "bullet", re.sub(r"^[-*]\s+", "", line).strip()
        else:
            yield "text", line.strip()


def _strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    return text


def export_to_docx(markdown_text: str, company_name: str) -> str:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    for kind, text in _parse_lines(markdown_text):
        text = _strip_md_inline(text)
        if kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=2)
        elif kind == "h3":
            doc.add_heading(text, level=3)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)

    filename = _safe_filename(company_name, "docx")
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath


class _ReportPDF(FPDF):
    def header(self):
        pass

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")


def export_to_pdf(markdown_text: str, company_name: str) -> str:
    pdf = _ReportPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    pdf.set_margins(18, 18, 18)

    for kind, text in _parse_lines(markdown_text):
        text = _strip_md_inline(text)
        if kind == "h1":
            pdf.set_font("Helvetica", "B", 18)
            pdf.multi_cell(0, 10, text)
            pdf.ln(2)
        elif kind == "h2":
            pdf.set_font("Helvetica", "B", 14)
            pdf.multi_cell(0, 9, text)
            pdf.ln(1)
        elif kind == "h3":
            pdf.set_font("Helvetica", "B", 12)
            pdf.multi_cell(0, 8, text)
        elif kind == "bullet":
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 7, f"-  {text}")
        else:
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 7, text)
        pdf.ln(1)

    filename = _safe_filename(company_name, "pdf")
    filepath = os.path.join(OUTPUT_DIR, filename)
    pdf.output(filepath)
    return filepath
